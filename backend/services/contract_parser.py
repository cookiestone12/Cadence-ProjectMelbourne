import os
import json
import logging
from typing import Dict, Any, Optional
from io import BytesIO

logger = logging.getLogger("cadence")

CONTRACT_TYPES = ["MASTER", "PUBLISHING", "SYNC_LICENSE", "DISTRIBUTION", "MECHANICAL", "PERFORMANCE", "OTHER"]
PAYMENT_DIRECTIONS = ["INCOMING", "OUTGOING"]
CURRENCIES = ["USD", "EUR", "GBP", "CAD", "AUD", "JPY"]


def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def parse_contract_document(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".pdf":
            text = extract_text_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            text = extract_text_from_docx(file_bytes)
        else:
            return {"success": False, "error": f"Unsupported file type: {ext}. Please upload a PDF or DOCX file."}
    except Exception as e:
        logger.error(f"Failed to extract text from {filename}: {e}")
        return {"success": False, "error": f"Could not read the document. Please ensure it is a valid {ext} file."}

    if not text or len(text.strip()) < 50:
        return {"success": False, "error": "The document appears to be empty or contains too little text to parse."}

    truncated_text = text[:12000]

    try:
        parsed = call_ai_parser(truncated_text)
    except Exception as e:
        logger.error(f"AI contract parsing failed: {e}")
        return {"success": False, "error": "AI analysis failed. Please try again or enter details manually."}

    text_preview = text[:500] + ("..." if len(text) > 500 else "")

    return {
        "success": True,
        "parsed_fields": parsed,
        "text_preview": text_preview,
    }


def call_ai_parser(document_text: str) -> Dict[str, Any]:
    from openai import OpenAI

    client = OpenAI(
        api_key=os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY"),
        base_url=os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")
    )

    prompt = f"""You are a music industry contract analyst. Analyze the following contract document text and extract structured information.

DOCUMENT TEXT:
---
{document_text}
---

Extract the following fields from the contract. Return a JSON object with these keys:

- "title": string — The name/title of the agreement (e.g., "Master Recording Agreement", "Publishing Deal Memo")
- "contract_type": string — One of: {json.dumps(CONTRACT_TYPES)}. Choose the best match based on the contract's subject matter.
- "payment_direction": string — "INCOMING" if the organization receives money, "OUTGOING" if the organization pays out. Default to "INCOMING" if unclear.
- "reference_number": string or null — Any reference/contract number mentioned
- "start_date": string or null — Effective/start date in YYYY-MM-DD format
- "end_date": string or null — Expiration/end date in YYYY-MM-DD format, or term end date
- "territory": array of strings — Territories mentioned (e.g., ["Worldwide"] or ["US", "Canada", "UK"])
- "advance_amount": number or null — Advance payment amount if mentioned (numeric value only, no currency symbols)
- "advance_currency": string — Currency code ({json.dumps(CURRENCIES)}). Default "USD" if not specified.
- "terms_summary": string — A concise 2-4 sentence summary of the key terms, obligations, and rights granted
- "notes": string or null — Any other notable provisions, restrictions, or special clauses worth highlighting
- "parties": array of objects — Each party involved, with keys:
  - "party_name": string — Name of the party
  - "party_role": string — Role (e.g., "Artist", "Publisher", "Label", "Producer", "Writer", "Licensor", "Licensee")
  - "contact_email": string or null — Email if found in the document

IMPORTANT:
1. Only extract information that is clearly stated in the document. Use null for fields you cannot determine.
2. For dates, convert to YYYY-MM-DD format. If only a year is given, use YYYY-01-01.
3. For territory, if the contract says "worldwide" or "world", return ["Worldwide"].
4. For advance_amount, extract only the numeric value (e.g., 50000 not "$50,000").
5. For terms_summary, focus on: rights granted, royalty rates/splits, term length, and key obligations.
6. Respond ONLY with the JSON object, no markdown formatting or other text."""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=2000
    )

    response_text = response.choices[0].message.content.strip()
    if response_text.startswith("```"):
        response_text = response_text.split("```")[1]
        if response_text.startswith("json"):
            response_text = response_text[4:]
        response_text = response_text.strip()

    parsed = json.loads(response_text)

    if parsed.get("contract_type") and parsed["contract_type"] not in CONTRACT_TYPES:
        parsed["contract_type"] = "OTHER"

    if parsed.get("payment_direction") and parsed["payment_direction"] not in PAYMENT_DIRECTIONS:
        parsed["payment_direction"] = "INCOMING"

    if parsed.get("advance_currency") and parsed["advance_currency"] not in CURRENCIES:
        parsed["advance_currency"] = "USD"

    if parsed.get("advance_amount"):
        try:
            parsed["advance_amount"] = float(parsed["advance_amount"])
        except (ValueError, TypeError):
            parsed["advance_amount"] = None

    if parsed.get("territory") and isinstance(parsed["territory"], str):
        parsed["territory"] = [parsed["territory"]]

    if not parsed.get("parties"):
        parsed["parties"] = []

    return parsed
